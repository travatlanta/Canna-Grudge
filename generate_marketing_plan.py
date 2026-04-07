from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# ── Color Palette ──────────────────────────────────────────
BLACK       = RGBColor(0x0D, 0x10, 0x17)
GOLD        = RGBColor(0xD4, 0xA8, 0x43)
GOLD_DARK   = RGBColor(0xA0, 0x7C, 0x28)
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_GRAY  = RGBColor(0xF5, 0xF5, 0xF5)
MID_GRAY    = RGBColor(0xE0, 0xE0, 0xE0)
DARK_GRAY   = RGBColor(0x3A, 0x3A, 0x3A)
SECTION_BG  = RGBColor(0x1A, 0x1A, 0x2E)

doc = Document()

# ── Page Margins ───────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Helper: set paragraph shading ──────────────────────────
def shade_paragraph(para, hex_color):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    pPr.append(shd)

def shade_cell(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        border = OxmlElement(f'w:{edge}')
        border.set(qn('w:val'), kwargs.get(edge, 'none'))
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), kwargs.get(f'{edge}_color', 'auto'))
        tcBorders.append(border)
    tcPr.append(tcBorders)

# ── Helper: add a horizontal rule ──────────────────────────
def add_hr(doc, color='D4A843'):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pb = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color)
    pb.append(bottom)
    pPr.append(pb)

# ── Helper: add cover-style heading block ──────────────────
def add_cover(doc):
    # Dark banner paragraph
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    shade_paragraph(p, '0D1017')
    run = p.add_run('🔥  CANNAGRUDGE')
    run.bold = True
    run.font.size = Pt(36)
    run.font.color.rgb = GOLD
    run.font.name = 'Calibri'

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after  = Pt(0)
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    shade_paragraph(p2, '0D1017')
    r2 = p2.add_run('OFFICIAL MARKETING PLAN  ·  APRIL 25, 2026')
    r2.bold = False
    r2.font.size = Pt(13)
    r2.font.color.rgb = RGBColor(0xC4, 0xCC, 0xDA)
    r2.font.name = 'Calibri'

    p3 = doc.add_paragraph()
    p3.paragraph_format.space_before = Pt(0)
    p3.paragraph_format.space_after  = Pt(8)
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    shade_paragraph(p3, '0D1017')
    r3 = p3.add_run("Dunn's Arena  ·  Litchfield Park, AZ  ·  21+")
    r3.font.size = Pt(11)
    r3.font.color.rgb = RGBColor(0x8E, 0x99, 0xAC)
    r3.font.name = 'Calibri'

# ── Helper: Section header ─────────────────────────────────
def add_section_header(doc, text):
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    shade_paragraph(p, '1A1A2E')
    run = p.add_run(f'  {text}')
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = GOLD
    run.font.name = 'Calibri'
    add_hr(doc)

# ── Helper: Phase header ───────────────────────────────────
def add_phase_header(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(2)
    shade_paragraph(p, 'D4A843')
    run = p.add_run(f'  {text}  ')
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0x0D, 0x10, 0x17)
    run.font.name = 'Calibri'

# ── Helper: Sub-heading ────────────────────────────────────
def add_subheading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = DARK_GRAY
    run.font.name = 'Calibri'
    # underline accent
    pPr = p._p.get_or_add_pPr()
    pb = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'D4A843')
    pb.append(bottom)
    pPr.append(pb)

# ── Helper: Body text ──────────────────────────────────────
def add_body(doc, text, bold=False, italic=False, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(11)
    run.font.color.rgb = color or DARK_GRAY
    run.font.name = 'Calibri'
    return p

# ── Helper: Bullet ─────────────────────────────────────────
def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.left_indent  = Inches(0.3)
    if bold_prefix:
        run1 = p.add_run(bold_prefix)
        run1.bold = True
        run1.font.size = Pt(11)
        run1.font.color.rgb = DARK_GRAY
        run1.font.name = 'Calibri'
        run2 = p.add_run(text)
        run2.font.size = Pt(11)
        run2.font.color.rgb = DARK_GRAY
        run2.font.name = 'Calibri'
    else:
        run = p.add_run(text)
        run.font.size = Pt(11)
        run.font.color.rgb = DARK_GRAY
        run.font.name = 'Calibri'

# ── Helper: Callout box ────────────────────────────────────
def add_callout(doc, text, bg='FFF8E7'):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.left_indent  = Inches(0.2)
    p.paragraph_format.right_indent = Inches(0.2)
    shade_paragraph(p, bg)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x5A, 0x40, 0x00)
    run.font.name = 'Calibri'

# ── Helper: Styled table ───────────────────────────────────
def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # Header row
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        shade_cell(cell, '0D1017')
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = GOLD
        run.font.name = 'Calibri'

    # Data rows
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        bg = 'F9F9F9' if r_idx % 2 == 0 else 'FFFFFF'
        for c_idx, cell_text in enumerate(row_data):
            cell = row.cells[c_idx]
            shade_cell(cell, bg)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if c_idx == 0 else WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(cell_text)
            run.font.size = Pt(10)
            run.font.name = 'Calibri'
            run.font.color.rgb = DARK_GRAY
            if c_idx == 0:
                run.bold = True

    # Column widths
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)

    doc.add_paragraph()

# ══════════════════════════════════════════════════════════════
#  BUILD DOCUMENT
# ══════════════════════════════════════════════════════════════

add_cover(doc)

# ─── SITUATION ────────────────────────────────────────────────
add_section_header(doc, 'THE SITUATION')
add_body(doc, 'April 25, 2026 — 46 days out. Here\'s what we\'re working with:')
add_bullet(doc, '3,000-person capacity event at Dunn\'s Arena, Litchfield Park, AZ — this is a major venue')
add_bullet(doc, 'Own ticketing platform on cannagrudge.com — no third-party cut')
add_bullet(doc, 'Large warm email database ready to deploy immediately')
add_bullet(doc, 'Daily dispensary presence across the Phoenix metro market')
add_bullet(doc, 'Nirvana partnership — 10 locations across the West Valley')
add_bullet(doc, '13 recognized AZ cannabis brands competing — each with their own audience')
add_bullet(doc, 'Dual audience: cannabis consumers + Phoenix combat sports fans')

# ─── AD SPEND ─────────────────────────────────────────────────
add_section_header(doc, 'PAID ADVERTISING BUDGET')
add_body(doc, 'Total recommended spend: ~$7,500–$9,000. At 3,000 capacity you need volume — not just warm email opens. This budget reflects a real push across every available channel. Concentrate 40% of Traffic Roots and Facebook budget in the final 14 days when purchase intent peaks.', italic=True)
doc.add_paragraph()

add_table(doc,
    headers=['Channel', 'Budget', 'What It Buys'],
    rows=[
        ['Traffic Roots — Retargeting',   '$600',   'Serve ads to email list + site visitors who didn\'t convert. Highest ROI — warm audience already exposed.'],
        ['Traffic Roots — Prospecting',   '$1,200', 'Cold geo-targeted display/mobile ads, Phoenix metro 21+ cannabis + nightlife audiences at scale'],
        ['Traffic Roots — Video Pre-Roll','$800',   '15-second video pre-roll on cannabis publisher sites — highest recall format, critical for a 3k event'],
        ['Facebook / Instagram Ads',      '$1,500', 'Frame as entertainment event. Target 21–45, Phoenix metro + surrounding cities, combat sports + nightlife. Retarget all site visitors.'],
        ['Google Display / Search',       '$800',   '"Cannabis event Phoenix," "fight night AZ," "things to do Phoenix April 2026" — high buyer intent searches'],
        ['Reddit Promoted Posts',         '$300',   'r/AZTrees, r/phoenix, r/arizona — promoted posts look native, no cannabis ad policy issues'],
        ['TikTok Ads',                    '$500',   'Spark Ads from brand partner posts that perform organically — amplify what\'s already working'],
        ['Influencer / Creator Outreach', '$500',   '2–4 AZ cannabis/nightlife creators posting about the event to their audiences. Negotiate tickets + cash.'],
        ['Flier Printing',                '$400',   '5,000 fliers with QR codes + promo codes for full dispensary market coverage'],
        ['Digital Screen Assets',         '$150',   'Formatted slides for in-store digital screens across dispensary network'],
        ['Budtender Incentive Prize Pool','$300',   'VIP tickets + product + cash — top referring locations compete hard when the stakes are real'],
        ['TOTAL',                         '~$7,050–$8,550', 'Full multi-channel push scaled for 3,000 seats'],
    ],
    col_widths=[2.5, 1.0, 3.5]
)

# ─── PHASE 1 ──────────────────────────────────────────────────
add_phase_header(doc, 'PHASE 1 — IGNITION  (March 10–20)')

add_subheading(doc, 'Email Blast #1 — "The Announcement"')
add_body(doc, 'Send this week. The database is warm — use it now.')
doc.add_paragraph()
add_body(doc, 'Subject Line (A/B test two):', bold=True)
add_callout(doc, '🔥 AZ\'s biggest cannabis event is 46 days away')
add_callout(doc, 'The brands you smoke every day are about to fight')
doc.add_paragraph()
add_body(doc, 'Email Strategy:', bold=True)
add_bullet(doc, 'Hero image — best event flier or graphic, full width')
add_bullet(doc, 'One punchy paragraph: what it is, where, when — no walls of text')
add_bullet(doc, 'List ALL 13 competing brand names — customers recognize them, instant credibility')
    add_bullet(doc, 'Highlight the experience: live boxing, 6 brand bouts, free tattoos, vendor games with trophies, vendor expo, music')
add_bullet(doc, 'One CTA button: "Get Tickets → cannagrudge.com" — nothing else')
add_bullet(doc, 'If it doesn\'t fit on a phone screen above the fold, it\'s too long')
add_bullet(doc, '', bold_prefix='Exclusive promo code: ')
add_callout(doc, 'EARLYBIRD15 — 15% off, expires March 24. Email list only. Makes them feel VIP.')
doc.add_paragraph()
add_body(doc, 'Send Specs:', bold=True)
add_bullet(doc, 'Tuesday or Wednesday, 10–11am AZ time')
add_bullet(doc, 'Segment: cannabis-adjacent contacts first, then full blast')
add_bullet(doc, 'Track open rate and CTR — this data tells you how warm the list actually is')

add_subheading(doc, 'Google Business Profile — Event Listing')
add_body(doc, 'Do this today. Free. Highest-intent local traffic you can get.')
add_bullet(doc, 'Create or claim the CannaGrudge Google Business Profile')
add_bullet(doc, 'Add the April 25 event — shows up in Google Search and Maps for West Valley searches')
add_bullet(doc, 'Keywords: "cannabis event Phoenix," "fight night Litchfield Park," "AZ brand battle"')
add_bullet(doc, 'Add event image, direct ticket link, full description')
add_bullet(doc, 'This is organic search placement you cannot buy at any reasonable price')

add_subheading(doc, 'Facebook Event — Create & Seed')
add_bullet(doc, 'Create an official Facebook Event for April 25 from the CannaGrudge page')
add_bullet(doc, 'Personal invites to every contact — personal invites convert higher than mass invites')
add_bullet(doc, 'Ask each competing brand to co-host or share the event from their page')
add_body(doc, 'Share immediately in these groups:', bold=True)
add_bullet(doc, 'Arizona Cannabis Community')
add_bullet(doc, 'AZ Stoners / Phoenix 420 groups')
add_bullet(doc, 'West Valley Phoenix events groups')
add_bullet(doc, 'Arizona Fight Fans / Combat Sports AZ groups — this is a boxing event too')
add_callout(doc, '💡 KEY INSIGHT: You have a dual audience — cannabis consumers AND boxing/MMA fans in AZ. Most marketing only targets one. A "celebrity boxing + cannabis" angle hits both — tap the fight fan community actively.')

add_subheading(doc, 'Reddit — r/AZTrees (50k+ Members)')
add_bullet(doc, 'Post: "CannaGrudge is April 25 — 13 AZ brands boxing for bragging rights at Dunn\'s Arena"')
add_bullet(doc, 'Include all brand names — fans of those brands will upvote and share')
add_bullet(doc, 'Reply to every comment in the first 2 hours (Reddit algorithm rewards early engagement)')
add_bullet(doc, '$20–$50 Reddit promoted post targeting r/AZTrees + r/phoenix for extra reach')
add_bullet(doc, 'Post countdown updates at 2-week and 1-week marks')

# ─── PHASE 2 ──────────────────────────────────────────────────
add_phase_header(doc, 'PHASE 2 — BUILD THE FIRE  (March 20 – April 7)')

add_subheading(doc, 'Email Blast #2 — "Meet the Brands"  (Send ~March 20–22)')
add_body(doc, 'Subject: ', bold=True)
add_callout(doc, 'The lineup is set. 13 brands. One winner. 🥊')
add_bullet(doc, 'Show all 13 brand logos — frame it like a fight card: "In the ring on April 25..."')
add_bullet(doc, 'One-liner on each brand — make their fans feel represented')
add_bullet(doc, 'Restate the full event experience')
add_bullet(doc, 'Show ticket tiers and pricing clearly')
add_bullet(doc, 'Remind about EARLYBIRD15 code with deadline pressure if still active')
add_callout(doc, '💡 Each brand has customers who will forward this email. You\'re leveraging 13 brands\' audiences in a single send — no cost.')

add_subheading(doc, 'Brand Partner Content — Activate Their Audiences')
add_body(doc, 'This is your biggest free marketing lever. Every competing brand has followers who are your exact target demo.')
doc.add_paragraph()
add_body(doc, 'What to give each brand:', bold=True)
add_bullet(doc, 'A co-branded graphic — their logo + CannaGrudge + April 25 date')
add_bullet(doc, '2–3 pre-written caption options for Instagram and Facebook (make it easy for them)')
add_bullet(doc, 'The direct ticket link')
add_bullet(doc, 'Ask for minimum one post per week between now and the event')
doc.add_paragraph()
add_body(doc, 'How to frame it to them:', bold=True)
add_callout(doc, '"Their fans vote with ticket purchases. If your brand gets the biggest turnout, you own those bragging rights. Push hard." — Give them a reason to care.')

add_subheading(doc, 'Traffic Roots — Launch Paid Display Campaign')
add_bullet(doc, 'Platform: trafficroots.com — the only fully cannabis-legal programmatic ad network')
add_bullet(doc, '', bold_prefix='Geo-target: ')
add_body(doc, '             Phoenix metro + specifically West Valley: Peoria, Glendale, Goodyear, Avondale, Buckeye, Litchfield Park')
add_bullet(doc, '', bold_prefix='Audience: ')
add_body(doc, '             21+, cannabis interest, nightlife + event attendees')
add_bullet(doc, '', bold_prefix='Formats: ')
add_body(doc, '             Static banner (300x250, 728x90) + mobile interstitial + video pre-roll')
add_bullet(doc, '', bold_prefix='Retargeting: ')
add_body(doc, '             Upload email list as custom audience — serve ads to openers who didn\'t convert. Highest-efficiency spend.')
add_bullet(doc, '', bold_prefix='Creative: ')
add_body(doc, '             Bold event graphic, date and venue visible at a glance, single GET TICKETS CTA')

add_subheading(doc, 'Instagram / TikTok Content Calendar')
add_table(doc,
    headers=['Day', 'Content Type', 'What to Post'],
    rows=[
        ['Monday',    'Brand Spotlight',    'One competing brand — hype their entry, tag them'],
        ['Wednesday', 'Behind the Scenes',  'Venue, prep, merch, team activity'],
        ['Thursday',  'Countdown',          '"30 days / 14 days until CannaGrudge" graphic'],
        ['Friday',    'Ticket Push',        '"Weekend plans? April 25 →" with ticket link'],
        ['Daily Story','Engagement',        'Countdown sticker, polls ("Who wins — Nirvana or Diablo?"), ticket link'],
    ],
    col_widths=[1.3, 1.8, 4.0]
)
add_callout(doc, '🎯 TikTok angle: "13 cannabis brands are about to settle who\'s #1 in AZ." Get brand reps to film 10-second trash-talk clips at each other. Pure engagement bait — this format can go viral in the cannabis content space.')

add_subheading(doc, 'Email Blast #3 — "The Experience"  (Send ~April 1)')
add_body(doc, 'Subject: ', bold=True)
add_callout(doc, 'Free tattoos. Live fights. Open bar vendors. April 25.')
add_body(doc, 'Shift the message from "what is it" to "what will you experience that night."')
add_bullet(doc, 'Walk them through the night hour by hour — make them feel like they\'re already there')
add_bullet(doc, 'Free tattoos — if confirmed, this is a massive hook for this demographic')
    add_bullet(doc, 'Vendor games with trophy finals, music/DJ, vendor expo — paint the full picture')
add_bullet(doc, 'Ticket tiers clearly explained — what each tier includes, price, availability')
add_bullet(doc, 'Add a sold ticket count or "X tickets remaining" if numbers support FOMO')

# ─── PHASE 3 ──────────────────────────────────────────────────
add_phase_header(doc, 'PHASE 3 — CLOSE THE SEATS  (April 7–22)')

add_subheading(doc, 'Email Blast #4 — "Two Weeks Out"  (Send April 7)')
add_body(doc, 'Subject: ', bold=True)
add_callout(doc, '2 weeks. 🔥 Last chance for the good seats.')
add_bullet(doc, 'Urgency: ticket count, capacity reminder (3,000-person venue — big enough to feel possible, real enough to sell out)')
add_bullet(doc, 'Show what they\'re missing — brand content, venue photos, hype material')
add_bullet(doc, '', bold_prefix='Introduce "bring a friend" incentive: ')
add_callout(doc, 'PLUS1 promo code — buy 2, second ticket 10% off. Drives group purchases.')

add_subheading(doc, 'Email Blast #5 — "Final 7 Days"  (Send April 18)')
add_body(doc, 'Subject: ', bold=True)
add_callout(doc, 'Saturday. You in or out?')
add_body(doc, 'Keep it extremely short. 3 sentences above the fold. Urgency through simplicity.')
add_bullet(doc, '"3,000 people. One night. Saturday April 25. You in?"')
add_bullet(doc, 'Large ticket button — only CTA on the email')
add_bullet(doc, 'Nothing else. The brevity is the message.')

add_subheading(doc, 'SMS Blast (If You Have Numbers)')
add_callout(doc, '📱 SMS open rate is 95%+ vs ~25% for email. If you have numbers, use them.')
add_bullet(doc, 'Text 1 — 7 days out: "CannaGrudge is THIS Saturday at Dunn\'s Arena. Tickets → [link]"')
add_bullet(doc, 'Text 2 — 24 hours out: "Tomorrow night. See you there 🔥 [link]"')
add_bullet(doc, 'Short, direct, no subject line needed — just the message and link')

add_subheading(doc, 'Free Local Event Listings')
add_body(doc, 'Low effort, free reach. Submit to all three today:')
add_bullet(doc, '', bold_prefix='Patch.com ')
add_body(doc, '             Litchfield Park / Goodyear / Peoria editions — free community event submissions')
add_bullet(doc, '', bold_prefix='AZCentral.com ')
add_body(doc, '             Events calendar — free submit, reaches large Phoenix metro audience actively looking for things to do')
add_bullet(doc, '', bold_prefix='Do602.com ')
add_body(doc, '             Phoenix events discovery site — cannabis-adjacent audience already using the platform')

# ─── DISPENSARY ───────────────────────────────────────────────
add_section_header(doc, 'DISPENSARY PRESENCE — MAXIMIZING DAILY FLOOR TIME')

add_body(doc, 'Your teams are already on the floor every day across a massive footprint. The goal is to stop treating these visits as flier drops and start running them as mini campaign activations. Every single visit should convert floor time into ticket sales.')

add_subheading(doc, 'Location Tier System')
add_table(doc,
    headers=['Tier', 'Locations', 'Visit Frequency', 'Activation Level'],
    rows=[
        ['Tier 1 — Anchor',  'Nirvana (all 10 locations)',        '2x per week', 'Full activation — script, screen, POS, budtender briefing'],
        ['Tier 2 — Partner', 'Mint + other partner accounts',     '1x per week', 'Standard activation — POS placement, budtender touchpoint'],
        ['Tier 3 — Coverage','Broader market dispensary accounts','1x per 2 weeks','Flier refresh + quick budtender mention'],
    ],
    col_widths=[1.4, 2.2, 1.4, 2.5]
)

add_subheading(doc, 'The Floor Activation Script')
add_body(doc, 'Train every team member to deliver this exact message. Consistency is everything:')
add_callout(doc, '"Hey — you know the brands on your shelf right here? [hold up flier] They\'re all boxing each other for bragging rights on April 25 at Dunn\'s Arena. 3,000 people, live fights, free tattoos, vendor expo — it\'s 21+. Tickets are $40, there\'s a promo code on the flier for a discount. It\'s going to sell out."')
doc.add_paragraph()
add_body(doc, 'Why this script works:', bold=True)
add_bullet(doc, 'Leads with brands they already recognize and trust — instant relevance')
add_bullet(doc, 'Gives them a physical flier with QR code — something to hold and scan later')
add_bullet(doc, 'Mentions 3,000 people — makes it sound like a real event, not a small local gathering')
add_bullet(doc, 'The tattoo detail is unusual enough to stick in memory')

add_subheading(doc, 'Vendor Day Checklist — Every Tier 1 & Tier 2 Visit')
add_body(doc, 'ON ARRIVAL:', bold=True)
add_bullet(doc, 'Restock fliers at POS — top of the stack, face up, QR code visible')
add_bullet(doc, 'Ask manager or lead budtender: "Can we get the CannaGrudge slide up on the screen this week?"')
add_bullet(doc, 'Drop a small stack in the waiting area')
doc.add_paragraph()
add_body(doc, 'DURING FLOOR TIME:', bold=True)
add_bullet(doc, 'Personally hand a flier to anyone browsing — do not leave it to them to pick one up')
add_bullet(doc, 'Deliver the activation script to as many customers as possible')
add_bullet(doc, 'Scan your own QR code in front of the customer to show it works and where it goes')
doc.add_paragraph()
add_body(doc, 'BEFORE LEAVING:', bold=True)
add_bullet(doc, 'Remind the on-duty budtender about the promo code contest — top store wins')
add_bullet(doc, 'Log the visit: location name, fliers left, screen updated Y/N, budtender conversation Y/N')

add_subheading(doc, 'Nirvana Partnership — 10 Locations, Maximum Activation')
add_body(doc, 'Nirvana is not just a distribution point — they are a competing brand with loyal customers who are your warmest audience. Use every channel available through them:')
add_bullet(doc, 'Ask Nirvana to blast the event to their customer loyalty database — a "Nirvana presents CannaGrudge" co-branded send')
add_bullet(doc, 'Get the event posted on Nirvana\'s official social channels — their followers already trust the brand')
add_bullet(doc, 'Push for 18x24 poster or counter display at all 10 locations — elevates beyond a flier')
add_bullet(doc, 'Get a CannaGrudge slide on rotation on all in-store digital screens')
add_callout(doc, '🏆 Nirvana angle for in-store: "Come cheer for us April 25" — their customers feel tribal loyalty. Nirvana fans showing up is brand support, not just ticket sales.')

add_subheading(doc, 'Budtender Incentive Contest')
add_body(doc, 'You have 15–20+ locations with multiple budtenders each. This is a free sales team — you just need to activate them.')
doc.add_paragraph()
add_body(doc, 'How it works:', bold=True)
add_bullet(doc, 'Every location gets a unique promo code tied to that store (e.g., NIRVANA75TH, NIRVANAPEORIA, MINTTEMPE)')
add_bullet(doc, 'Codes are already trackable in your promo code dashboard')
add_bullet(doc, 'Top 3 locations by redemptions each win: 2 VIP tickets + branded CannaGrudge package')
add_bullet(doc, 'Announce the contest to every manager today so budtenders know from day one')
add_bullet(doc, 'Send a mid-contest leaderboard update around April 10 — "Location X is in the lead" — drives competition')
add_callout(doc, '💰 A $300 prize pool activating 20+ locations can generate 100–200+ ticket sales across your network. Real prizes drive real competition — this is the highest ROI item in the entire plan.')

# ─── FLIER CHECKLIST ──────────────────────────────────────────
add_section_header(doc, 'FLIER REQUIREMENTS CHECKLIST')
add_body(doc, 'Every flier hitting dispensary floors must have all of the following. A flier without a QR code and promo code is not working hard enough:')
add_bullet(doc, 'Event name + date + venue — readable in under 2 seconds')
add_bullet(doc, 'QR code linking directly to cannagrudge.com/tickets')
add_bullet(doc, 'A promo code for a discount (use DISPENSARY for general floor fliers)')
add_bullet(doc, 'All 13 brand logos — customers recognize them, it\'s instant social proof')
add_bullet(doc, '21+ notice — required')
add_bullet(doc, 'cannagrudge.com URL')

# ─── PROMO CODES ──────────────────────────────────────────────
add_section_header(doc, 'PROMO CODE STRATEGY')
add_body(doc, 'All codes are live in your existing promo code system. Deploy them by channel so you can track exactly which sources are converting:')
doc.add_paragraph()
add_table(doc,
    headers=['Code', 'Discount', 'Who Gets It', 'Expires'],
    rows=[
        ['EARLYBIRD15',    '15% off',               'Email Blast #1 — subscribers only',          'March 24'],
        ['DISPENSARY',     '$5 off',                 'General fliers on all dispensary floors',     'April 20'],
        ['NIRVANA[STORE]', '$5 off',                 'Store-specific codes for budtender contest',  'April 20'],
        ['PLUS1',          '10% off 2+ tickets',     'Email Blast #4 — bring a friend',             'April 22'],
        ['LASTCALL',       '$3 off',                 'Final email blast + social media',            'April 24'],
    ],
    col_widths=[1.8, 1.3, 2.8, 1.2]
)

# ─── METRICS ──────────────────────────────────────────────────
add_section_header(doc, 'METRICS — WHAT TO TRACK EVERY MONDAY')
add_body(doc, 'Your analytics dashboard is already built at cannagrudge.com/analytics. Check these weekly:')
add_bullet(doc, '', bold_prefix='Ticket page visits vs. conversions — ')
add_body(doc, '             If traffic is up but conversions are flat, the page or price is the issue, not the marketing.')
add_bullet(doc, '', bold_prefix='Promo code redemptions by code — ')
add_body(doc, '             Tells you exactly which channel is actually selling tickets. Follow the money.')
add_bullet(doc, '', bold_prefix='Email open rates — ')
add_body(doc, '             Below 20% means subject lines need work. Above 35% means the list is hot — send more.')
add_bullet(doc, '', bold_prefix='Traffic sources — ')
add_body(doc, '             Which referrers are sending buyers. Double down on what\'s working.')

# ─── THE ONE THING ────────────────────────────────────────────
add_section_header(doc, 'THE ONE THING')
add_callout(doc, '🥊 If you only do one thing beyond sending the email blast — get every competing brand to post about CannaGrudge to their own Instagram this week.\n\n13 brands. Each with thousands of followers. Every follower is your exact target demo. They have a personal reason to promote because their brand is competing.\n\nThat\'s earned media you cannot buy at any price.', bg='FFF3CD')

# ─── TIMELINE SUMMARY ─────────────────────────────────────────
add_section_header(doc, 'EVENT TIMELINE SUMMARY')
add_table(doc,
    headers=['Week', 'Dates', 'Priority Actions'],
    rows=[
        ['Week 1',  'Mar 10–16', 'Email Blast #1 · Google Business listing · Facebook Event · Budtender contest launch · r/AZTrees post'],
        ['Week 2',  'Mar 17–23', 'Brand partner content activation · Traffic Roots campaign launch · Daily social posting starts'],
        ['Week 3',  'Mar 24–30', 'Email Blast #2 (brand roster) · Facebook/Instagram ads live · Nirvana loyalty email blast'],
        ['Week 4',  'Mar 31–6',  'Email Blast #3 (the experience) · Mid-contest leaderboard update · Ramp social posting'],
        ['Week 5',  'Apr 7–13',  'Email Blast #4 (2 weeks out) · PLUS1 code live · Heavy paid ad spend begins'],
        ['Week 6',  'Apr 14–20', 'Final dispensary push · Free local listings (Patch, AZCentral, Do602)'],
        ['Week 7',  'Apr 21–24', 'Email Blast #5 · SMS blast (if available) · LASTCALL code · Full social blitz'],
        ['Event Day','Apr 25',   'Behind-the-scenes content · Live posts · Brand arrival coverage · Social throughout the night'],
    ],
    col_widths=[0.8, 1.1, 5.3]
)

# ─── FOOTER ───────────────────────────────────────────────────
doc.add_paragraph()
add_hr(doc)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('CannaGrudge  ·  April 25, 2026  ·  Dunn\'s Arena, Litchfield Park AZ  ·  cannagrudge.com')
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x8E, 0x99, 0xAC)
run.font.name = 'Calibri'

# ── Save ──────────────────────────────────────────────────────
output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'CannaGrudge_Marketing_Plan.docx')
doc.save(output_path)
print(f'Saved: {output_path}')
